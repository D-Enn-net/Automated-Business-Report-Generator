# report_builder.py
from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF
from datetime import datetime
import os
from config import OUTPUT_REPORTS_PATH

def create_word_report(metrics_data, chart_path):
    """
    Generates a comprehensive Word document with analysis, tables, and visualizations.
    """
    print("\nGenerating comprehensive Word report...")
    doc = Document()
    today_date = datetime.now().strftime("%Y-%m-%d")

    # Title
    doc.add_heading(f"Sales Report - {today_date}", level=1)
    
    # Executive Summary
    total_sales = metrics_data['total_sales']
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(f"This report provides an analysis of the Superstore sales data. ")
    doc.add_paragraph(f"The total sales figure across all regions and categories is ${total_sales:,.2f}.")
    doc.add_paragraph("")

    # Sales by Region Section
    doc.add_heading("Sales Performance by Region", level=2)
    doc.add_paragraph("The bar chart below illustrates the sales distribution across different regions.")
    doc.add_picture(chart_path, width=Inches(6.0))
    doc.add_paragraph("The table below provides the detailed sales figures for each region.")
    
    # Add a data table for regions
    region_data = metrics_data['sales_by_region'].reset_index() # Convert Series to DataFrame
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Region'
    hdr_cells[1].text = 'Total Sales ($)'
    for index, row in region_data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Region'])
        row_cells[1].text = f"{row['Sales']:,.2f}"
    doc.add_paragraph("")

    # Sales by Category Section
    doc.add_heading("Sales Performance by Category", level=2)
    category_data = metrics_data['sales_by_category'].reset_index() # Convert Series to DataFrame
    doc.add_paragraph("The table below provides the detailed sales figures for each product category.")
    
    # Add a data table for categories
    table_cat = doc.add_table(rows=1, cols=2)
    table_cat.style = 'Table Grid'
    hdr_cells_cat = table_cat.rows[0].cells
    hdr_cells_cat[0].text = 'Category'
    hdr_cells_cat[1].text = 'Total Sales ($)'
    for index, row in category_data.iterrows():
        row_cells_cat = table_cat.add_row().cells
        row_cells_cat[0].text = str(row['Category'])
        row_cells_cat[1].text = f"{row['Sales']:,.2f}"
        
    # Save Document
    report_path = os.path.join(OUTPUT_REPORTS_PATH, f'Sales_Report_{today_date}.docx')
    doc.save(report_path)
    print(f"Word report saved successfully to: {report_path}")
    return report_path

def create_pdf_report(metrics_data, chart_path):
    """
    Generates a final, consistent, and complete PDF document with all tables.
    """
    print("\nGenerating final, high-quality PDF report...")
    pdf = FPDF()
    pdf.add_page()
    
    # Title
    today_date = datetime.now().strftime("%Y-%m-%d")
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, f"Sales Report - {today_date}", ln=True, align='C')
    pdf.ln(10)
    
    # Executive Summary
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, "Executive Summary", ln=True)
    pdf.set_font('Arial', '', 12)
    total_sales = metrics_data['total_sales']
    pdf.multi_cell(0, 10, f"This report provides an analysis of the Superstore sales data. The total sales figure across all regions and categories is ${total_sales:,.2f}.")
    pdf.ln(5)
    
    # Sales by Region Section (Chart & Table) ---
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, "Sales Performance by Region", ln=True)
    pdf.image(chart_path, w=180)
    pdf.ln(5)

    # Region Table
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(0, 8, "The table below provides the detailed sales figures for each region.")
    pdf.ln(2)
    
    pdf.set_font('Arial', 'B', 12)
    col_width_region = 100
    col_width_val = 90
    pdf.cell(col_width_region, 10, 'Region', border=1)
    pdf.cell(col_width_val, 10, 'Total Sales ($)', border=1, ln=1)
    
    pdf.set_font('Arial', '', 12)
    region_data = metrics_data['sales_by_region'].reset_index()
    for index, row in region_data.iterrows():
        pdf.cell(col_width_region, 10, str(row['Region']), border=1)
        pdf.cell(col_width_val, 10, f"{row['Sales']:,.2f}", border=1, ln=1, align='R')
    pdf.ln(5)

    # Sales by Category Section
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, "Sales Performance by Category", ln=True)
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(0, 8, "The table below provides the detailed sales figures for each product category.")
    pdf.ln(2)
    
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(col_width_region, 10, 'Category', border=1)
    pdf.cell(col_width_val, 10, 'Total Sales ($)', border=1, ln=1)
    
    pdf.set_font('Arial', '', 12)
    category_data = metrics_data['sales_by_category'].reset_index()
    for index, row in category_data.iterrows():
        pdf.cell(col_width_region, 10, str(row['Category']), border=1)
        pdf.cell(col_width_val, 10, f"{row['Sales']:,.2f}", border=1, ln=1, align='R')
    
    # Save Document
    report_path = os.path.join(OUTPUT_REPORTS_PATH, f'Sales_Report_{today_date}.pdf')
    pdf.output(report_path)
    
    print(f"PDF report saved successfully to: {report_path}")
    return report_path